"""
File Parser Module — PDF, DOCX, TXT Text Extraction
=====================================================

📚 WHAT THIS MODULE DOES:
    Extracts plain text from uploaded files (PDF, DOCX, TXT).
    Users upload a file → this module reads it → returns clean text string.

📚 DESIGN PATTERN: Strategy Pattern
    ──────────────────────────────────
    The Strategy pattern defines a family of algorithms (parsing strategies),
    encapsulates each one, and makes them interchangeable.

    WHY?
    - PDF, DOCX, and TXT all need different parsing logic
    - But the caller (app.py) shouldn't care HOW the parsing works
    - It just calls: parser.parse(file) → gets text back

    STRUCTURE:
        FileParser (Abstract Base Class)  ← defines the interface
            ├── PDFParser                 ← strategy for .pdf
            ├── DOCXParser                ← strategy for .docx
            └── TXTParser                 ← strategy for .txt

        ParserFactory.get_parser(ext)     ← picks the right strategy

    IN AN INTERVIEW:
        "I used the Strategy pattern so that adding a new file format
        (e.g., .epub, .html) only requires adding one new class —
        no changes to existing code. This follows the Open/Closed Principle."

📚 LIBRARIES USED:
    - abc (Abstract Base Class): Built-in Python module for defining interfaces
    - PyPDF2: Reads PDF files page-by-page
    - pdfplumber: More accurate PDF extraction (handles tables, columns)
    - python-docx: Reads DOCX (Word) files paragraph-by-paragraph
    - chardet: Auto-detects text file encoding
    - pathlib: Modern filesystem path handling
    - io.BytesIO: Treats bytes in memory as a file-like object

📚 KEY CONCEPTS:
    - PDF: NOT a plain text format. It's a Page Description Language that stores
      drawing commands ("put character 'H' at position x=72, y=700"). PyPDF2
      reverses this to reconstruct the text.

    - DOCX: Actually a ZIP archive containing XML files. The main content is in
      word/document.xml. python-docx parses this XML for you.

    - Encoding: Text files can use different byte-to-character mappings:
      UTF-8 (most common, supports all languages), Latin-1 (Western European),
      ASCII (English only, 7-bit). Opening a file with the wrong encoding
      produces garbled text (mojibake). chardet detects the encoding automatically.
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional
import io
import logging

import chardet
import PyPDF2
import pdfplumber
from docx import Document

# Create a logger for this module
# 📚 Each module gets its own logger with __name__ (e.g., "utils.file_parser").
#    This lets you filter logs by module: "Show me only file_parser logs."
logger = logging.getLogger(__name__)


# ══════════════════════════════════════════════
# ABSTRACT BASE CLASS (Interface)
# ══════════════════════════════════════════════

class FileParser(ABC):
    """
    Abstract base class defining the interface for all file parsers.

    📚 WHAT IS AN ABSTRACT BASE CLASS (ABC)?
        - It's a class that CANNOT be instantiated directly
        - It defines methods that subclasses MUST implement
        - Think of it as a "contract" or "blueprint"

        Example:
            parser = FileParser()  # ❌ TypeError! Can't instantiate abstract class
            parser = PDFParser()   # ✅ Concrete subclass, allowed

    📚 WHY USE ABC?
        - Forces all parsers to have the same interface (parse method)
        - If you forget to implement parse() in a subclass → instant error
        - The caller doesn't need to know which parser it's using

    📚 @abstractmethod:
        Decorator that marks a method as "must be overridden by subclasses".
        If a subclass doesn't implement it → TypeError at instantiation time.
    """

    @abstractmethod
    def parse(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from file content.

        Args:
            file_content: Raw bytes of the uploaded file.
                📚 WHY bytes, not str?
                    Files from Streamlit's uploader arrive as raw bytes.
                    We don't know the encoding yet (for text files), and
                    binary files (PDF, DOCX) aren't text at all.

            filename: Original filename (used for logging and error messages).

        Returns:
            Extracted text as a string.

        Raises:
            FileParsingError: If the file cannot be parsed.
        """
        ...  # Ellipsis means "to be implemented by subclasses"

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        """File extensions this parser handles (e.g., ('.pdf',))."""
        return ()


# ══════════════════════════════════════════════
# CUSTOM EXCEPTION
# ══════════════════════════════════════════════

class FileParsingError(Exception):
    """
    Custom exception for file parsing failures.

    📚 WHY CUSTOM EXCEPTIONS?
        Instead of raising generic Exception("PDF failed"), we raise
        FileParsingError("PDF failed"). This lets callers handle parsing
        errors differently from other errors:

        try:
            text = parser.parse(data, "file.pdf")
        except FileParsingError as e:
            st.error(f"Could not read your file: {e}")
        except Exception as e:
            st.error(f"Unexpected error: {e}")
    """

    def __init__(self, filename: str, reason: str) -> None:
        self.filename = filename
        self.reason = reason
        super().__init__(f"Failed to parse '{filename}': {reason}")


# ══════════════════════════════════════════════
# PDF PARSER
# ══════════════════════════════════════════════

class PDFParser(FileParser):
    """
    Extracts text from PDF files.

    📚 HOW PDF TEXT EXTRACTION WORKS:
        PDFs store "draw character 'A' at position (x, y)" commands.
        Extraction algorithms must:
        1. Read each page's content stream
        2. Group characters by position into words
        3. Group words by position into lines
        4. Combine lines into paragraphs

        This is why PDF extraction is imperfect — the "text" isn't stored
        as text, it's stored as drawing instructions.

    📚 TWO LIBRARIES, WHY?
        - PyPDF2: Fast, lightweight, works for most simple PDFs
        - pdfplumber: More accurate for complex layouts (tables, columns)
        We try PyPDF2 first (faster), fall back to pdfplumber if it fails.

    📚 io.BytesIO:
        Wraps raw bytes in a file-like object. Libraries like PyPDF2 expect
        a "file" they can .read() and .seek() — BytesIO provides this
        interface without writing to disk.
    """

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return (".pdf",)

    def parse(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from a PDF file.

        Strategy:
            1. Try PyPDF2 (fast, simple)
            2. If PyPDF2 yields little/no text, try pdfplumber (more robust)
            3. If both fail, raise FileParsingError

        Args:
            file_content: Raw PDF bytes.
            filename: Original filename for logging.

        Returns:
            Extracted text as a single string.
        """
        logger.info("Parsing PDF: %s (%d bytes)", filename, len(file_content))

        # Attempt 1: PyPDF2
        text = self._parse_with_pypdf2(file_content, filename)

        # If PyPDF2 yielded very little text, try pdfplumber
        if len(text.strip()) < 50:
            logger.info("PyPDF2 extracted minimal text, trying pdfplumber...")
            pdfplumber_text = self._parse_with_pdfplumber(file_content, filename)
            if len(pdfplumber_text.strip()) > len(text.strip()):
                text = pdfplumber_text

        if not text.strip():
            raise FileParsingError(
                filename,
                "No text could be extracted. The PDF might be image-based "
                "(scanned). OCR is required for scanned PDFs."
            )

        logger.info("PDF parsed successfully: %d characters extracted", len(text))
        return text

    def _parse_with_pypdf2(self, file_content: bytes, filename: str) -> str:
        """
        Extract text using PyPDF2.

        📚 PyPDF2.PdfReader:
            - Reads the PDF structure (pages, metadata, fonts)
            - .pages: List of page objects
            - .extract_text(): Returns text from a single page
        """
        try:
            # BytesIO wraps bytes as a file-like object
            pdf_file = io.BytesIO(file_content)
            reader = PyPDF2.PdfReader(pdf_file)

            pages_text: list[str] = []
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text)
                    logger.debug("Page %d: %d chars", page_num + 1, len(page_text))

            return "\n\n".join(pages_text)

        except Exception as e:
            logger.warning("PyPDF2 failed for %s: %s", filename, e)
            return ""

    def _parse_with_pdfplumber(self, file_content: bytes, filename: str) -> str:
        """
        Extract text using pdfplumber (more robust for complex layouts).

        📚 pdfplumber vs PyPDF2:
            pdfplumber uses a more sophisticated layout analysis algorithm.
            It understands columns, tables, and text positioning better.
            Trade-off: It's slower and uses more memory.
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pages_text: list[str] = []

            with pdfplumber.open(pdf_file) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                        logger.debug(
                            "pdfplumber page %d: %d chars",
                            page_num + 1, len(page_text),
                        )

            return "\n\n".join(pages_text)

        except Exception as e:
            logger.warning("pdfplumber failed for %s: %s", filename, e)
            return ""


# ══════════════════════════════════════════════
# DOCX PARSER
# ══════════════════════════════════════════════

class DOCXParser(FileParser):
    """
    Extracts text from Microsoft Word (.docx) files.

    📚 HOW DOCX FILES WORK:
        A .docx file is actually a ZIP archive containing XML files:
            word/document.xml    ← main content (paragraphs, text runs)
            word/styles.xml      ← formatting styles
            word/media/          ← embedded images
            [Content_Types].xml  ← file type registry

        python-docx parses word/document.xml and gives us:
        - doc.paragraphs: List of Paragraph objects
        - Each paragraph has .text (the plain text content)
        - Each paragraph has .runs (text segments with formatting)

    📚 WHAT IS A "RUN"?
        A paragraph like "Hello **bold** world" has 3 runs:
        - Run 1: "Hello " (normal)
        - Run 2: "bold" (bold formatting)
        - Run 3: " world" (normal)
        paragraph.text concatenates all runs → "Hello bold world"
    """

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return (".docx",)

    def parse(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from a DOCX file.

        Extracts all paragraphs and joins them with newlines.
        Empty paragraphs (blank lines in the document) are preserved
        as paragraph separators.

        Args:
            file_content: Raw DOCX bytes.
            filename: Original filename for logging.

        Returns:
            Extracted text as a single string.
        """
        logger.info("Parsing DOCX: %s (%d bytes)", filename, len(file_content))

        try:
            # Document() can accept a file-like object
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)

            # Extract text from each paragraph
            paragraphs: list[str] = []
            for para in doc.paragraphs:
                # para.text gives the concatenated text of all runs
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            if not paragraphs:
                raise FileParsingError(
                    filename,
                    "No text found in the DOCX file. "
                    "The document might contain only images or tables."
                )

            result = "\n\n".join(paragraphs)
            logger.info(
                "DOCX parsed successfully: %d paragraphs, %d characters",
                len(paragraphs), len(result),
            )
            return result

        except FileParsingError:
            raise  # Re-raise our custom exception as-is
        except Exception as e:
            raise FileParsingError(filename, str(e)) from e


# ══════════════════════════════════════════════
# TXT PARSER
# ══════════════════════════════════════════════

class TXTParser(FileParser):
    """
    Extracts text from plain text (.txt) files with encoding detection.

    📚 CHARACTER ENCODING — A FUNDAMENTAL CONCEPT:
        Computers store text as numbers. Encoding is the mapping:
        - ASCII: 'A' = 65, 'B' = 66 (English only, 128 chars)
        - Latin-1: Extends ASCII to 256 chars (Western European: é, ñ, ü)
        - UTF-8: Variable-length encoding supporting ALL characters (emoji, 中文, العربية)
                  Backward-compatible with ASCII.

        If you open a Latin-1 file as UTF-8, you get "mojibake":
            "café" → "cafÃ©" (garbled text)

        chardet examines byte patterns to guess the encoding:
        - UTF-8 has distinctive byte patterns (110xxxxx 10xxxxxx for 2-byte chars)
        - Latin-1 has bytes in the 0x80-0xFF range
        - chardet uses statistical models trained on many languages

    📚 WHY NOT ALWAYS USE UTF-8?
        Many legacy systems, older documents, and Windows Notepad files
        use different encodings. We must handle whatever the user uploads.
    """

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return (".txt",)

    def parse(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from a TXT file with automatic encoding detection.

        Strategy:
            1. Try UTF-8 first (most common modern encoding)
            2. If that fails, use chardet to detect the encoding
            3. If chardet can't determine encoding, try Latin-1 (accepts any byte)

        Args:
            file_content: Raw bytes of the text file.
            filename: Original filename for logging.

        Returns:
            Decoded text as a string.
        """
        logger.info("Parsing TXT: %s (%d bytes)", filename, len(file_content))

        # Strategy 1: Try UTF-8 (most common)
        try:
            text = file_content.decode("utf-8")
            logger.info("TXT decoded as UTF-8: %d characters", len(text))
            return text
        except UnicodeDecodeError:
            logger.debug("UTF-8 decoding failed, trying chardet detection...")

        # Strategy 2: Auto-detect encoding with chardet
        detection = chardet.detect(file_content)
        detected_encoding: Optional[str] = detection.get("encoding")
        confidence: float = detection.get("confidence", 0.0)

        logger.info(
            "chardet detected encoding: %s (confidence: %.1f%%)",
            detected_encoding, confidence * 100,
        )

        if detected_encoding:
            try:
                text = file_content.decode(detected_encoding)
                logger.info(
                    "TXT decoded as %s: %d characters",
                    detected_encoding, len(text),
                )
                return text
            except (UnicodeDecodeError, LookupError) as e:
                logger.warning(
                    "Decoding with %s failed: %s", detected_encoding, e,
                )

        # Strategy 3: Fallback to Latin-1 (never fails — maps bytes 0-255 directly)
        logger.warning("Falling back to Latin-1 encoding for %s", filename)
        text = file_content.decode("latin-1")
        return text


# ══════════════════════════════════════════════
# PARSER FACTORY
# ══════════════════════════════════════════════

class ParserFactory:
    """
    Factory that returns the correct parser based on file extension.

    📚 DESIGN PATTERN: Factory Pattern
        Instead of writing:
            if ext == ".pdf":
                parser = PDFParser()
            elif ext == ".docx":
                parser = DOCXParser()
            ...

        We use a factory that encapsulates this logic:
            parser = ParserFactory.get_parser(".pdf")

        WHY?
        - Single Responsibility: Factory handles parser selection
        - Open/Closed: Add new formats by adding to the registry, not modifying code
        - Testability: Easy to mock the factory in tests

    📚 CLASS METHOD vs INSTANCE METHOD:
        @classmethod receives the class itself (cls) as the first argument,
        not an instance (self). You call it on the class directly:
            ParserFactory.get_parser(".pdf")  # No need to create an instance
    """

    # Registry mapping extensions to parser classes
    _parsers: dict[str, type[FileParser]] = {
        ".pdf": PDFParser,
        ".docx": DOCXParser,
        ".txt": TXTParser,
    }

    @classmethod
    def get_parser(cls, filename: str) -> FileParser:
        """
        Get the appropriate parser for a file based on its extension.

        Args:
            filename: The filename or path (e.g., "report.pdf").

        Returns:
            An instance of the appropriate FileParser subclass.

        Raises:
            ValueError: If the file extension is not supported.
        """
        # Path(filename).suffix returns the extension: ".pdf", ".docx", etc.
        ext = Path(filename).suffix.lower()

        parser_class = cls._parsers.get(ext)
        if parser_class is None:
            supported = ", ".join(cls._parsers.keys())
            raise ValueError(
                f"Unsupported file type: '{ext}'. "
                f"Supported types: {supported}"
            )

        logger.debug("Selected parser: %s for extension %s", parser_class.__name__, ext)
        return parser_class()

    @classmethod
    def supported_extensions(cls) -> list[str]:
        """Return list of all supported file extensions."""
        return list(cls._parsers.keys())

    @classmethod
    def register_parser(cls, extension: str, parser_class: type[FileParser]) -> None:
        """
        Register a new parser for a file extension.

        📚 WHY THIS METHOD?
            Makes the factory extensible. If you later add an EPUBParser:
                ParserFactory.register_parser(".epub", EPUBParser)
            No need to modify the factory's source code.

        Args:
            extension: File extension (e.g., ".epub").
            parser_class: The parser class to use for this extension.
        """
        cls._parsers[extension.lower()] = parser_class
        logger.info("Registered parser %s for extension %s", parser_class.__name__, extension)


# ══════════════════════════════════════════════
# CONVENIENCE FUNCTION
# ══════════════════════════════════════════════

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    High-level convenience function to extract text from any supported file.

    This is the function that app.py will call — simple, one-line interface.

    📚 WHY A CONVENIENCE FUNCTION?
        The caller doesn't need to know about parsers, factories, or strategies.
        It just calls: text = extract_text_from_file(uploaded_bytes, "report.pdf")
        All the complexity is hidden behind this clean interface.

    Args:
        file_content: Raw bytes of the file.
        filename: Original filename (determines which parser to use).

    Returns:
        Extracted text as a string.

    Raises:
        ValueError: If the file type is not supported.
        FileParsingError: If the file cannot be parsed.

    Example:
        >>> with open("report.pdf", "rb") as f:
        ...     text = extract_text_from_file(f.read(), "report.pdf")
        >>> print(text[:100])
        'Introduction: This report presents...'
    """
    parser = ParserFactory.get_parser(filename)
    return parser.parse(file_content, filename)
