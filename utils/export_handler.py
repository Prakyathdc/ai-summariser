"""
Export Handler Module
=======================
Handles exporting summaries into TXT, DOCX, and PDF formats.

📚 HOW FILE EXPORT WORKS:
    - TXT: Simple string encoding to bytes.
    - DOCX: Uses `python-docx` to construct a new document programmatically.
    - PDF: Uses `fpdf2` to construct a PDF by drawing text cell by cell.
"""

from io import BytesIO
from docx import Document
from fpdf import FPDF
import logging

logger = logging.getLogger(__name__)

class ExportHandler:
    @staticmethod
    def export_to_txt(summary: str, original: str) -> bytes:
        """Returns TXT file content as bytes."""
        content = f"--- AI SUMMARY ---\n{summary}\n\n--- ORIGINAL TEXT ---\n{original}"
        return content.encode('utf-8')

    @staticmethod
    def export_to_docx(summary: str, original: str) -> bytes:
        """Returns DOCX file content as bytes."""
        doc = Document()
        doc.add_heading('AI Summary Report', 0)
        
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary)
        
        doc.add_heading('Original Text', level=1)
        doc.add_paragraph(original)
        
        file_stream = BytesIO()
        doc.save(file_stream)
        return file_stream.getvalue()

    @staticmethod
    def export_to_pdf(summary: str, original: str) -> bytes:
        """Returns PDF file content as bytes."""
        pdf = FPDF()
        pdf.add_page()
        
        # We must use a built-in font for simplicity, or handle Unicode carefully.
        # FPDF's default font doesn't support complex unicode easily without a TTF font.
        # So we encode with latin-1 replace to prevent crashes with emoji/special chars.
        
        def safe_text(txt: str) -> str:
            return txt.encode('latin-1', 'replace').decode('latin-1')

        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "AI Summary Report", ln=True, align='C')
        pdf.ln(10)

        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Summary:", ln=True)
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 8, safe_text(summary))
        pdf.ln(10)

        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, "Original Text:", ln=True)
        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 8, safe_text(original))

        return pdf.output(dest='S')
